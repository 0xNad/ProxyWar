from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageEnhance, ImageFont, ImageOps
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_lets_play_proxywar_brief import (
    BLUE,
    BLUE_DARK,
    CALLOUT,
    GREEN,
    LINE,
    MUTED,
    NAVY,
    PALE,
    PALE_BLUE,
    SLATE,
    USABLE_WIDTH_DXA,
    WHITE,
    add_body,
    add_bullet,
    add_callout,
    add_code_block,
    add_hyperlink,
    add_page_break,
    add_page_number,
    add_table,
    add_three_cards,
    rgb,
    set_cell_border,
    set_cell_margins,
    set_cell_shading,
    set_cell_text,
    set_cell_width,
    set_paragraph_bottom_border,
    set_row_cant_split,
    set_run_font,
    set_table_geometry,
    style_document,
)


ROOT = Path("/Users/claude/Documents/proxywar_main")
WORK = ROOT / "tmp" / "lets_play_proxywar_player_guide"
OUT = ROOT / "outputs" / "share" / "Lets_Play_Proxy_War_New_Player_Guide.docx"
DRAFT = WORK / "Lets_Play_Proxy_War_New_Player_Guide_draft.docx"
SCREENSHOT = ROOT / "resources" / "images" / "GameplayScreenshot.png"
HERO = WORK / "hero.png"
ANNOTATED = WORK / "battlefield_annotated.png"
TURN_LOOP = WORK / "turn_loop.png"
ITERATION_LOOP = WORK / "iteration_loop.png"
CODING_AGENT_LOOP = WORK / "coding_agent_loop.png"

GOLD = "D69E2E"
PALE_GOLD = "FFF7E0"
TEAL = "197278"
PALE_TEAL = "E9F7F6"
CODE_BG = "F6F8FA"


def configure_compact_reference_styles(doc: Document) -> None:
    """Resolve the compact_reference_guide preset exactly.

    Named overrides: 31 pt editorial cover title, gold cover kicker, and
    player-notebook code/output cells.
    """
    style_document(doc)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, BLUE_DARK, 10, 5),
    ):
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.size = Pt(10.4)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    doc.styles["PW Code"].font.size = Pt(8.4)
    doc.styles["PW Code"].paragraph_format.line_spacing = 1.0
    doc.styles["PW Code"].paragraph_format.space_before = Pt(0)
    doc.styles["PW Code"].paragraph_format.space_after = Pt(0)


def configure_header_footer(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "PROXY WAR  /  PLAYER NOTEBOOK"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.runs[0], size=8.5, color=MUTED, bold=True)
    set_paragraph_bottom_border(paragraph, color=LINE)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    add_page_number(paragraph)

    section.different_first_page_header_footer = True
    first_header = section.first_page_header
    paragraph = first_header.paragraphs[0]
    paragraph.text = "PROXY WAR"
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.runs[0], size=8.5, color=BLUE, bold=True)
    set_paragraph_bottom_border(paragraph, color=BLUE, size=8)

    first_footer = section.first_page_footer
    paragraph = first_footer.paragraphs[0]
    paragraph.text = "PLAYER NOTEBOOK  •  JULY 2026"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    set_run_font(paragraph.runs[0], size=8.2, color=MUTED, bold=True)


def make_hero() -> None:
    source = Image.open(SCREENSHOT).convert("RGB")
    source = ImageEnhance.Contrast(source).enhance(1.08)
    source = ImageEnhance.Color(source).enhance(1.08)
    source = ImageOps.expand(source, border=8, fill=f"#{BLUE}")
    source.save(HERO, quality=95)


def make_annotated_battlefield() -> None:
    source = Image.open(SCREENSHOT).convert("RGB")
    source = ImageEnhance.Contrast(source).enhance(1.05)
    draw = ImageDraw.Draw(source, "RGBA")
    label_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 22)
    detail_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 18)
    callouts = [
        {
            "label": "TERRITORY",
            "detail": "controlled land",
            "box": (55, 345, 300, 420),
            "start": (300, 383),
            "target": (355, 392),
        },
        {
            "label": "ECONOMY",
            "detail": "city and factory",
            "box": (280, 70, 530, 145),
            "start": (405, 145),
            "target": (387, 235),
        },
        {
            "label": "FRONTIER",
            "detail": "two nations meet",
            "box": (440, 525, 710, 600),
            "start": (575, 525),
            "target": (555, 405),
        },
        {
            "label": "WATER",
            "detail": "sea and naval routes",
            "box": (930, 420, 1160, 495),
            "start": (930, 458),
            "target": (850, 505),
        },
        {
            "label": "NEUTRAL LAND",
            "detail": "unconquered island",
            "box": (930, 530, 1160, 605),
            "start": (930, 568),
            "target": (810, 575),
        },
    ]
    for callout in callouts:
        start = callout["start"]
        target = callout["target"]
        draw.line((start, target), fill=(255, 255, 255, 235), width=10)
        draw.line((start, target), fill=(11, 37, 69, 255), width=5)
        x, y = target
        draw.ellipse((x - 9, y - 9, x + 9, y + 9), fill=(255, 255, 255, 255), outline=(11, 37, 69, 255), width=4)
    for callout in callouts:
        x0, y0, x1, y1 = callout["box"]
        draw.rounded_rectangle((x0 + 4, y0 + 5, x1 + 4, y1 + 5), radius=12, fill=(11, 37, 69, 65))
        draw.rounded_rectangle((x0, y0, x1, y1), radius=12, fill=(255, 255, 255, 240), outline=(11, 37, 69, 255), width=3)
        draw.text((x0 + 14, y0 + 9), callout["label"], font=label_font, fill=(11, 37, 69, 255))
        draw.text((x0 + 14, y0 + 39), callout["detail"], font=detail_font, fill=(71, 85, 105, 255))
    source = ImageOps.expand(source, border=8, fill=f"#{BLUE}")
    source.save(ANNOTATED, quality=95)


def make_horizontal_flow(path: Path, nodes: list[tuple[str, str]], accent: str = BLUE) -> None:
    width, height = 1800, 330
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    font_number = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 34)
    font_title = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 34)
    font_small = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 27)
    count = len(nodes)
    gap = 58
    left = 24
    box_w = int((width - left * 2 - gap * (count - 1)) / count)
    box_h = 205
    top = 55
    for idx, (title, subtitle) in enumerate(nodes):
        x0 = left + idx * (box_w + gap)
        x1 = x0 + box_w
        y0, y1 = top, top + box_h
        fill = PALE_BLUE if idx % 2 == 0 else CALLOUT
        draw.rounded_rectangle((x0, y0, x1, y1), radius=18, fill=f"#{fill}", outline=f"#{accent}", width=4)
        draw.ellipse((x0 + 16, y0 + 16, x0 + 64, y0 + 64), fill=f"#{accent}")
        number = str(idx + 1)
        bbox = draw.textbbox((0, 0), number, font=font_number)
        draw.text((x0 + 40 - (bbox[2] - bbox[0]) / 2, y0 + 39 - (bbox[3] - bbox[1]) / 2 - 3), number, font=font_number, fill=f"#{WHITE}")
        wrapped = title.split("\n")
        y = y0 + 83
        for line in wrapped:
            bbox = draw.textbbox((0, 0), line, font=font_title)
            draw.text((x0 + (box_w - (bbox[2] - bbox[0])) / 2, y), line, font=font_title, fill=f"#{NAVY}")
            y += 39
        bbox = draw.textbbox((0, 0), subtitle, font=font_small)
        draw.text((x0 + (box_w - (bbox[2] - bbox[0])) / 2, y1 - 39), subtitle, font=font_small, fill=f"#{SLATE}")
        if idx < count - 1:
            arrow_y = y0 + box_h / 2
            start = x1 + 9
            end = x1 + gap - 9
            draw.line((start, arrow_y, end, arrow_y), fill=f"#{accent}", width=6)
            draw.polygon([(end, arrow_y), (end - 18, arrow_y - 12), (end - 18, arrow_y + 12)], fill=f"#{accent}")
    image.save(path, quality=95)


def make_diagrams() -> None:
    make_hero()
    make_annotated_battlefield()
    make_horizontal_flow(
        TURN_LOOP,
        [
            ("Read the\nstate", "observation"),
            ("Review legal\nactions", "offered moves"),
            ("Choose one\naction ID", "your policy"),
            ("See what\nhappened", "replay + result"),
        ],
    )
    make_horizontal_flow(
        ITERATION_LOOP,
        [
            ("Change one\nbehavior", "small edit"),
            ("Play a\nmatch", "same contract"),
            ("Watch the\nreplay", "visible outcome"),
            ("Inspect the\ndecisions", "reason + health"),
            ("Compare and\nrepeat", "next version"),
        ],
        accent=TEAL,
    )
    make_horizontal_flow(
        CODING_AGENT_LOOP,
        [
            ("Capture one\nmoment", "replay evidence"),
            ("State the\nhypothesis", "desired behavior"),
            ("Make one\npolicy edit", "bounded change"),
            ("Run matched\ntests", "comparable setup"),
            ("Verify and\ndecide", "keep or revert"),
        ],
        accent=GOLD,
    )


def add_picture(doc: Document, path: Path, alt: str, width: float = 6.35, title: str = "Proxy War") -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("title", title)
    shape._inline.docPr.set("descr", alt)


def add_kicker(doc: Document, text: str, color: str = BLUE) -> None:
    paragraph = doc.add_paragraph(style="PW Kicker")
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=9, color=color, bold=True)


def add_code_cell(doc: Document, label: str, code: str, *, output: str | None = None) -> None:
    paragraph = doc.add_paragraph(style="PW Label")
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(label.upper())
    set_run_font(run, size=8.2, color=BLUE, bold=True)
    add_code_block(doc, code)
    if output is not None:
        table = doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [USABLE_WIDTH_DXA])
        cell = table.cell(0, 0)
        set_cell_shading(cell, PALE_TEAL)
        set_cell_margins(cell, top=95, start=120, bottom=95, end=120)
        set_cell_border(
            cell,
            start={"val": "single", "sz": "12", "color": TEAL},
            top={"val": "nil"},
            bottom={"val": "nil"},
            end={"val": "nil"},
        )
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(output)
        set_run_font(run, name="Consolas", size=8.7, color=NAVY)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_metric_strip(doc: Document, metrics: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [USABLE_WIDTH_DXA // len(metrics)] * len(metrics)
    widths[-1] += USABLE_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths, indent=110)
    for idx, (value, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, PALE_BLUE if idx % 2 == 0 else CALLOUT)
        set_cell_margins(cell, top=120, start=110, bottom=120, end=110)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": LINE},
            bottom={"val": "single", "sz": "4", "color": LINE},
            start={"val": "single", "sz": "4", "color": LINE},
            end={"val": "single", "sz": "4", "color": LINE},
        )
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(value)
        set_run_font(run, size=17, color=BLUE_DARK, bold=True)
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label.upper())
        set_run_font(run, size=7.8, color=MUTED, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def create_decimal_numbering(doc: Document) -> int:
    """Create a real, independently restarting decimal list definition."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 0
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_list(doc: Document, items: list[str | tuple[str, str]]) -> None:
    num_id = create_decimal_numbering(doc)
    for item in items:
        text, bold_lead = (item, None) if isinstance(item, str) else item
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        p_pr.append(num_pr)
        if bold_lead and text.startswith(bold_lead):
            run = paragraph.add_run(bold_lead)
            set_run_font(run, size=10.2, bold=True)
            run = paragraph.add_run(text[len(bold_lead):])
            set_run_font(run, size=10.2)
        else:
            run = paragraph.add_run(text)
            set_run_font(run, size=10.2)


def add_decision_receipt(doc: Document) -> None:
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    widths = [2250, 7110]
    set_table_geometry(table, widths)
    values = [
        ("Moment", "Spawn turn"),
        ("Chosen action", "spawn:631446"),
        ("Reason", "Explore a strong opening position."),
        ("Decision status", "Accepted"),
        ("Effect", "Confirmed: the nation spawned, remained alive, and owned territory."),
        ("Next question", "Did the opening position create room to expand and build?"),
    ]
    for row, (label, value) in zip(table.rows, values):
        set_row_cant_split(row)
        left, right = row.cells
        set_cell_width(left, widths[0])
        set_cell_width(right, widths[1])
        set_cell_margins(left, top=95, bottom=95)
        set_cell_margins(right, top=95, bottom=95)
        set_cell_shading(left, PALE)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_text(left, label, bold=True, size=9.0)
        set_cell_text(right, value, size=9.0)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def remove_trailing_empty_paragraph(doc: Document) -> None:
    if not doc.paragraphs:
        return
    paragraph = doc.paragraphs[-1]
    if paragraph.text.strip():
        return
    paragraph._element.getparent().remove(paragraph._element)


def build() -> Path:
    WORK.mkdir(parents=True, exist_ok=True)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    make_diagrams()

    doc = Document()
    configure_compact_reference_styles(doc)
    configure_header_footer(doc.sections[0])

    properties = doc.core_properties
    properties.title = "Let's Play Proxy War"
    properties.subject = "A hands-on first agent guide"
    properties.author = "Proxy War"
    properties.keywords = "Proxy War, Coworld, autonomous agent, replay, LegalAction"
    properties.comments = "Player-facing tutorial."

    # Page 1 — editorial cover.
    add_kicker(doc, "Proxy War  •  Player notebook", GOLD)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run("Let’s Play Proxy War")
    set_run_font(run, size=31, color=NAVY, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run("Every nation on the map is controlled by code. Yours starts with one function.")
    set_run_font(run, size=14, color=SLATE, italic=True)

    add_picture(
        doc,
        HERO,
        "A Proxy War match showing colored territories, borders, structures, ports, and rival nations across Europe.",
        width=6.35,
        title="Proxy War battlefield",
    )

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run("BUILD AN AUTONOMOUS NATION  •  PLAY A MATCH  •  LEARN FROM THE REPLAY")
    set_run_font(run, size=9.3, color=BLUE, bold=True)

    add_three_cards(
        doc,
        [
            ("READ THE TURN", "Understand the observation and the moves available right now."),
            ("CHANGE THE POLICY", "Edit one strategy function and choose an exact offered action."),
            ("LEARN FROM PLAY", "Watch the replay, inspect the decisions, and improve one behavior."),
        ],
    )

    add_page_break(doc)

    # Page 2 — game orientation.
    add_kicker(doc, "1  •  Meet the battlefield")
    doc.add_heading("Every nation is an agent", level=1)
    add_body(
        doc,
        "Proxy War is a simultaneous-turn territorial strategy game for autonomous players. Your policy chooses where to spawn, when to expand, what to build, which rivals to pressure, when to cooperate, and when to wait. The map turns those choices into a visible story.",
        keep=True,
    )
    add_picture(
        doc,
        ANNOTATED,
        "Annotated Proxy War battlefield. Labeled callouts point directly to territory, an economy structure, a frontier, open water, and neutral land on an unconquered island.",
        width=6.35,
        title="Meet the battlefield",
    )
    add_table(
        doc,
        ("Map feature", "Why it matters"),
        [
            ("Territory", "Controlled land creates options, resources, and pressure."),
            ("Frontier", "Where two nations meet, attack, defense, alliance, and retreat decisions become urgent."),
            ("Economy", "Cities, factories, and upgrades turn land into long-term strength."),
            ("Water", "Sea separates land and creates naval routes for ports, transports, and warships."),
            ("Neutral land", "Unconquered land can be expanded into without attacking another nation."),
        ],
        (1800, 7560),
    )
    add_callout(
        doc,
        "Your objective",
        "Control the map. Expand into open land, build an economy, survive pressure, and convert a strong position into the best result.",
        fill=PALE_GOLD,
        title_color=GOLD,
    )

    add_page_break(doc)

    # Page 3 — one turn and the exact-ID contract.
    add_kicker(doc, "2  •  Understand one turn")
    doc.add_heading("The game offers choices. Your policy picks one.", level=1)
    add_body(
        doc,
        "At each decision, your player receives a structured observation plus the actions that are legal at that moment. It returns one action ID and a short reason. The map advances, and the result becomes part of the replay.",
        keep=True,
    )
    add_picture(
        doc,
        TURN_LOOP,
        "Four-step player loop: read the state, review legal actions, choose one action ID, and see what happened in the replay and result.",
        width=6.35,
        title="One Proxy War turn",
    )
    add_code_cell(
        doc,
        "Incoming turn — shortened",
        """const message = {
  type: \"decision_request\",
  requestID: \"req_example\",
  request: {
    observation: {
      phase: \"active\", turnNumber: 400,
      ownState: { tilesOwned: 52, troops: 62518, gold: \"209800\" }
    },
    legalActions: [
      { id: \"expand:terra-nullius:10\", kind: \"attack\", risk: { level: \"low\" }, metadata: { expansion: true } },
      { id: \"expand:terra-nullius:20\", kind: \"attack\", risk: { level: \"low\" }, metadata: { expansion: true } },
      { id: \"expand:terra-nullius:35\", kind: \"attack\", risk: { level: \"medium\" }, metadata: { expansion: true } },
      { id: \"build:City:132564\", kind: \"build\", risk: { level: \"low\" } },
      { id: \"hold\", kind: \"hold\", risk: { level: \"none\" } }
    ]
  }
};
const { observation, legalActions } = message.request;""",
    )
    add_code_cell(
        doc,
        "Your response",
        """const response = {
  type: \"decision_response\",
  requestID: message.requestID,
  selectedLegalActionId: \"expand:terra-nullius:10\",
  reason: \"Early neutral expansion is safe and useful.\",
  confidence: 0.84
};""",
    )
    add_callout(
        doc,
        "The one rule",
        "selectedLegalActionId must exactly match one id from the offered legalActions list. Your policy chooses the move; Proxy War validates it and sends it through the game path.",
        fill=PALE_TEAL,
        title_color=TEAL,
    )

    add_page_break(doc)

    # Page 4 — the starter policy.
    add_kicker(doc, "3  •  Build your starter")
    doc.add_heading("Strategy lives in one function", level=1)
    add_body(
        doc,
        "Open coworld-adapter/src/starter-player.mjs and replace its chooseAction() function with this rule policy. On page 6, you will also update the call so the function receives the current observation.",
        keep=True,
    )
    add_code_cell(
        doc,
        "Policy cell",
        """function chooseAction(actions, observation = {}) {
  if (!Array.isArray(actions) || actions.length === 0) {
    throw new Error(\"decision_request had no legalActions\");
  }

  const safe = actions.filter(action => action.risk?.level !== \"high\");
  const pick = test => safe.find(test);

  if (observation.phase === \"spawn\") {
    return pick(action => action.kind === \"spawn\") ?? actions[0];
  }

  return (
    pick(action => action.kind === \"attack\" && action.metadata?.expansion === true) ??
    pick(action => action.kind === \"build\") ??
    pick(action => action.kind === \"alliance_request\") ??
    pick(action => action.kind === \"upgrade_structure\") ??
    actions.find(action => action.kind === \"hold\") ??
    actions[0]
  );
}""",
    )
    doc.add_heading("Read it in plain English", level=2)
    add_bullet(doc, "Spawn when the match is waiting for a starting position.", bold_lead="Spawn")
    add_bullet(doc, "Expand into neutral land while a non-high-risk expansion is offered.", bold_lead="Expand")
    add_bullet(doc, "Build, form an alliance, or upgrade when expansion is unavailable.", bold_lead="Build")
    add_bullet(doc, "Hold when no higher-priority move remains.", bold_lead="Hold")
    add_code_cell(
        doc,
        "Run against the turn from section 2",
        """const action = chooseAction(legalActions, observation);
console.log(action.id);""",
        output="expand:terra-nullius:10",
    )
    add_callout(
        doc,
        "Checkpoint",
        "You now have a policy that reads the current menu, prefers non-high-risk moves, and returns one offered action ID every turn.",
        fill=PALE_BLUE,
    )

    add_page_break(doc)

    # Page 5 — one controlled change.
    add_kicker(doc, "4  •  Make your first strategy change")
    doc.add_heading("Commit more troops to the opening", level=1)
    add_body(
        doc,
        "The baseline takes the first neutral-expansion action, which is the 10% option in our frozen turn. Make one deliberate change: while your nation owns fewer than 100 tiles, prefer the offered 20% expansion.",
        keep=True,
    )
    add_code_cell(
        doc,
        "Insert before the default return",
        """if ((observation.ownState?.tilesOwned ?? 0) < 100) {
  const strongerOpening = pick(
    action => action.id === \"expand:terra-nullius:20\"
  );
  if (strongerOpening) return strongerOpening;
}""",
        output="selectedLegalActionId  →  expand:terra-nullius:20",
    )
    add_three_cards(
        doc,
        [
            ("BEFORE", "The starter selects the first offered neutral expansion: 10%."),
            ("CHANGE", "During the opening, prefer the offered 20% commitment."),
            ("AFTER", "The same turn now produces expand:terra-nullius:20."),
        ],
    )
    doc.add_heading("What to look for in the replay", level=2)
    add_numbered_list(
        doc,
        [
            "Did the stronger opening claim useful land faster?",
            "Did the nation keep enough troops to survive pressure?",
            "Did the rule stop applying once the opening was over?",
        ],
    )
    add_callout(
        doc,
        "The experiment",
        "The protocol stayed the same. You changed one behavior, then created a visible before-and-after question for the next match.",
        fill=PALE_GOLD,
        title_color=GOLD,
    )
    doc.add_heading("Three next experiments", level=2)
    add_bullet(doc, "Economy: prefer a City after the first land grab.", bold_lead="Economy:")
    add_bullet(doc, "Defense: hold or retreat when incoming pressure rises.", bold_lead="Defense:")
    add_bullet(doc, "Diplomacy: request an alliance when cooperation creates room to grow.", bold_lead="Diplomacy:")

    add_page_break(doc)

    # Page 6 — connect the policy and run an episode.
    add_kicker(doc, "5  •  Play a match")
    doc.add_heading("Connect your policy to the episode", level=1)
    add_body(
        doc,
        "The starter already handles the websocket connection and reply. Change its action call so chooseAction() receives both the offered actions and the current observation.",
        keep=True,
    )
    add_code_cell(
        doc,
        "Update this call in starter-player.mjs",
        """const action = chooseAction(
  message.request.legalActions ?? [],
  message.request.observation ?? {}
);""",
    )
    add_callout(
        doc,
        "Your player is complete",
        "Keep the starter transport, replace chooseAction() with the function from page 4, and make the call change above.",
        fill=PALE_BLUE,
    )
    doc.add_heading("Your first run", level=2)
    add_body(
        doc,
        "Have Docker running, Node 24+ installed, and uv available. Then run this from the Proxy War repository:",
        keep=True,
    )
    add_code_cell(
        doc,
        "Build and play",
        """cd coworld-adapter
PROXYWAR_REPO=.. npm run build:image
docker tag proxywar-coworld-local:latest proxywar-coworld-local:coworld-3e7e218fc73f
npm run run:episode""",
    )
    add_numbered_list(
        doc,
        [
            "Save coworld-adapter/src/starter-player.mjs with both edits.",
            "Run the four commands above.",
            "Open coworld-adapter/coworld/results/results.json, then watch the replay from the beginning.",
        ],
    )
    add_callout(
        doc,
        "What success looks like",
        "The player connects, returns offered action IDs, keeps acting after spawn, and produces a result plus a watchable replay.",
        fill=PALE_TEAL,
        title_color=TEAL,
    )

    add_page_break(doc)

    # Page 7 — replay payoff.
    add_kicker(doc, "6  •  Watch the replay")
    doc.add_heading("The match is the test. The replay is the lesson.", level=1)
    add_body(
        doc,
        "Start with the visible story: where the nation spawned, how quickly it expanded, when it built, who it allied with, and where its position stalled or broke through. Then open the decision trail for the moments that need explanation.",
        keep=True,
    )
    doc.add_heading("Example result", level=2)
    add_metric_strip(
        doc,
        [
            ("0.510", "score"),
            ("367", "tiles owned"),
            ("28 / 28", "accepted"),
            ("0", "fallbacks"),
        ],
    )
    add_body(
        doc,
        "This episode ended with both players alive. The leading nation held 367 tiles against 353, so its normalized territory score was 0.510.",
        keep=True,
    )
    doc.add_heading("One decision receipt", level=2)
    add_decision_receipt(doc)
    doc.add_heading("Read the evidence in three passes", level=2)
    add_numbered_list(
        doc,
        [
            ("Map: Did the nation expand, build, cooperate, and apply pressure at sensible moments?", "Map:"),
            ("Decision: Did the chosen action and reason match the state it saw?", "Decision:"),
            ("Health: Did fallbacks, degraded decisions, or repeated actions take control away from the policy?", "Health:"),
        ],
    )
    add_callout(
        doc,
        "Read acceptance correctly",
        "Accepted means the choice passed the decision-validation path. Use the effect audit and replay to see what changed and whether it was strategically useful.",
        fill=PALE_GOLD,
        title_color=GOLD,
    )

    add_page_break(doc)

    # Page 8 — turn replay evidence into a bounded coding-agent task.
    add_kicker(doc, "7  •  Work with your coding agent")
    doc.add_heading("Turn replay evidence into one testable change", level=1)
    add_body(
        doc,
        "A coding agent is most useful when it receives one specific failure, the evidence behind it, and a narrow definition of success. Give it one experiment - not a vague request to make the player smarter.",
        keep=True,
    )
    add_picture(
        doc,
        CODING_AGENT_LOOP,
        "Five-step coding-agent experiment loop: capture one replay moment, state a hypothesis, make one bounded policy edit, run matched tests, then verify and decide whether to keep or revert.",
        width=6.35,
        title="Coding-agent experiment loop",
    )
    doc.add_heading("Give your coding agent this task", level=2)
    add_callout(
        doc,
        "Copy and complete",
        "At [turn or replay moment], the player selected [offered action ID]. Its recorded reason was [reason]. The replay showed [effect], but I expected [desired behavior]. Change only [policy rule]. Keep the starter connection code unchanged and select only IDs offered on that turn. Save the current version as the baseline. Run comparable baseline and changed episodes. Report whether the rule activated; the replay and decision evidence; score, tiles, accepted, fallback, and degraded counts; and keep or revert against [success criterion].",
        fill=PALE_GOLD,
        title_color=GOLD,
    )
    doc.add_heading("Judge the change", level=2)
    add_bullet(doc, "Baseline: save the current version and result; write the success criterion before editing.", bold_lead="Baseline:")
    add_bullet(doc, "Match: use the same map, opponent mix, and starting settings whenever the runner allows it; run several episodes.", bold_lead="Match:")
    add_bullet(doc, "Verify: confirm in the decision receipts that the intended rule actually activated.", bold_lead="Verify:")
    add_bullet(doc, "Decide: keep only if the written criterion is met without new fallbacks or degraded decisions; otherwise revise or revert.", bold_lead="Decide:")
    doc.add_heading("When the result surprises you", level=2)
    add_three_cards(
        doc,
        [
            ("NO CONNECTION", "Fix the run setup or starter transport before judging strategy."),
            ("REJECTED / FALLBACK", "Check that the returned ID exactly matched an offered legal action."),
            ("MATCH UNCHANGED", "The new rule may not have activated. Inspect the decision receipts."),
        ],
    )

    add_page_break(doc)

    # Page 9 — iteration and close.
    add_kicker(doc, "8  •  Improve your player")
    doc.add_heading("Change one behavior. Play again.", level=1)
    add_body(
        doc,
        "A replay turns every match into a focused next step. Keep the contract fixed, change one policy choice, and compare the visible outcome plus the decision evidence.",
        keep=True,
    )
    add_picture(
        doc,
        ITERATION_LOOP,
        "Five-step improvement loop: change one behavior, play a match, watch the replay, inspect the decisions, compare and repeat.",
        width=6.35,
        title="Proxy War improvement loop",
    )
    doc.add_heading("Your next four experiments", level=2)
    add_numbered_list(
        doc,
        [
            ("Opening commitment: compare 10%, 20%, and 35% neutral expansion.", "Opening commitment:"),
            ("Economy build order: compare an early City with an early Port while expansion continues.", "Economy build order:"),
            ("Diplomacy under pressure: request help before a stronger neighbor attacks.", "Diplomacy under pressure:"),
            ("Anti-stall memory: stop repeating an action that has stopped producing value.", "Anti-stall memory:"),
        ],
    )
    doc.add_heading("Your first scorecard", level=2)
    add_bullet(doc, "The nation acted after spawn and kept making non-hold decisions.")
    add_bullet(doc, "Each reason explained the choice in terms of the current state.")
    add_bullet(doc, "The replay showed a recognizable strategy rather than random legal moves.")
    add_bullet(doc, "Fallback and degraded counts stayed visible in the result.")
    add_bullet(doc, "The next version changes one behavior and preserves everything else.")
    add_callout(
        doc,
        "You are ready",
        "A strong first agent starts simple: legal choices, continuous play, and clear evidence for the next revision. Build, play, watch, inspect, improve.",
        fill=PALE_TEAL,
        title_color=TEAL,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Continue with the Proxy War Coworld starter  →  ")
    set_run_font(run, size=9.5, color=MUTED, bold=True)
    add_hyperlink(
        paragraph,
        "github.com/0xNad/ProxyWar/tree/main/coworld-adapter",
        "https://github.com/0xNad/ProxyWar/tree/main/coworld-adapter",
    )
    remove_trailing_empty_paragraph(doc)

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(DRAFT)
    return DRAFT


if __name__ == "__main__":
    print(build())
