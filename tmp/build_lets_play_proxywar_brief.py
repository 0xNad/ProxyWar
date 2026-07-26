from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/claude/Documents/proxywar_main")
WORK = ROOT / "tmp" / "lets_play_proxywar_brief"
OUT_DIR = ROOT / "outputs" / "share"
DRAFT = WORK / "Lets_Play_Proxy_War_Collaboration_Brief_draft.docx"
DIAGRAM = WORK / "proxywar_decision_flow.png"

PAGE_WIDTH_DXA = 12240
USABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

NAVY = "0B2545"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
SLATE = "475569"
MUTED = "667085"
PALE = "E8EEF5"
PALE_BLUE = "EEF5FB"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
LINE = "C9D5E3"
GREEN = "2F6B55"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs) -> None:
    """Set individual cell borders. kwargs keys: top, bottom, start, end, insideH, insideV."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, spec in kwargs.items():
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in spec.items():
            element.set(qn(f"w:{key}"), str(value))


def set_row_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths: Sequence[int], indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != USABLE_WIDTH_DXA:
        raise ValueError(f"Column widths must total {USABLE_WIDTH_DXA}; got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        set_row_cant_split(row)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_headers(table) -> None:
    if table.rows:
        set_row_repeat(table.rows[0])


def set_cell_text(cell, text: str, bold: bool = False, color: str = NAVY, size: float = 9.3) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)


def set_run_font(run, name: str = "Calibri", size: float = 10.5, color: str = NAVY, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE, underline: bool = True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), color)
    r_pr.append(r_color)
    r_sz = OxmlElement("w:sz")
    r_sz.set(qn("w:val"), "19")
    r_pr.append(r_sz)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("PAGE ")
    set_run_font(prefix, size=8.5, color=MUTED)
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(cached)
    run._r.append(fld_char2)


def set_paragraph_bottom_border(paragraph, color: str = LINE, size: int = 6, space: int = 5) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 12),
        ("Subtitle", 13.5, SLATE, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, BLUE_DARK, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name not in ("Subtitle",)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for list_style_name in ("List Bullet", "List Number"):
        style = doc.styles[list_style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(10.2)
        style.font.color.rgb = rgb(NAVY)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    styles = doc.styles
    code = styles.add_style("PW Code", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = rgb(NAVY)
    code.paragraph_format.left_indent = Inches(0.16)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    code.paragraph_format.keep_together = True

    kicker = styles.add_style("PW Kicker", 1)
    kicker.font.name = "Calibri"
    kicker._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    kicker.font.size = Pt(9)
    kicker.font.bold = True
    kicker.font.color.rgb = rgb(BLUE)
    kicker.paragraph_format.space_after = Pt(8)

    label = styles.add_style("PW Label", 1)
    label.font.name = "Calibri"
    label._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    label.font.size = Pt(8)
    label.font.bold = True
    label.font.color.rgb = rgb(BLUE)
    label.paragraph_format.space_after = Pt(3)
    label.paragraph_format.keep_with_next = True

    small = styles.add_style("PW Small", 1)
    small.font.name = "Calibri"
    small._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    small.font.size = Pt(8.8)
    small.font.color.rgb = rgb(MUTED)
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.12


def configure_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.text = "PROXY WAR  /  GUIDED AGENT LAB"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.runs[0]
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    set_paragraph_bottom_border(p)

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    add_page_number(p)


def set_first_page_header_footer(section) -> None:
    section.different_first_page_header_footer = True
    first_header = section.first_page_header
    p = first_header.paragraphs[0]
    p.text = "PROXY WAR"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.runs[0]
    set_run_font(run, size=8.5, color=BLUE, bold=True)
    set_paragraph_bottom_border(p, color=BLUE, size=8)
    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.text = "COLLABORATION BRIEF  •  JULY 2026"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    set_run_font(p.runs[0], size=8.2, color=MUTED, bold=True)


def add_rich_paragraph(doc: Document, parts: Iterable[tuple[str, dict]], style: str | None = None, after: float | None = None, keep: bool = False):
    p = doc.add_paragraph(style=style)
    for text, attrs in parts:
        run = p.add_run(text)
        set_run_font(
            run,
            name=attrs.get("name", "Calibri"),
            size=attrs.get("size", 10.5),
            color=attrs.get("color", NAVY),
            bold=attrs.get("bold", False),
            italic=attrs.get("italic", False),
        )
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = keep
    return p


def add_body(doc: Document, text: str, *, bold_lead: str | None = None, after: float | None = None, keep: bool = False):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        r = p.add_run(text)
        set_run_font(r)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = keep
    return p


def add_label(doc: Document, text: str, proposed: bool = False):
    p = doc.add_paragraph(style="PW Label")
    run = p.add_run(("PROPOSED" if proposed else "REPOSITORY-VERIFIED") + "  /  " + text.upper())
    set_run_font(run, size=8, color=(GREEN if not proposed else BLUE), bold=True)
    return p


def add_callout(doc: Document, title: str, text: str, fill: str = CALLOUT, title_color: str = BLUE, icon: str | None = None):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [USABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=120, bottom=130, end=120)
    set_cell_border(
        cell,
        start={"val": "single", "sz": "16", "color": title_color},
        top={"val": "nil"},
        bottom={"val": "nil"},
        end={"val": "nil"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    heading = (f"{icon}  " if icon else "") + title
    r = p.add_run(heading)
    set_run_font(r, size=10.5, color=title_color, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_code_block(doc: Document, code_text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [USABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "4", "color": LINE},
        bottom={"val": "single", "sz": "4", "color": LINE},
        start={"val": "single", "sz": "4", "color": LINE},
        end={"val": "single", "sz": "4", "color": LINE},
    )
    p = cell.paragraphs[0]
    p.style = "PW Code"
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code_text)
    set_run_font(run, name="Consolas", size=8.3, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc: Document, text: str, *, bold_lead: str | None = None, style: str = "List Bullet"):
    p = doc.add_paragraph(style=style)
    spacer = " " if style == "List Bullet" else ""
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(spacer + bold_lead)
        set_run_font(r, size=10.2, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, size=10.2)
    else:
        r = p.add_run(spacer + text)
        set_run_font(r, size=10.2)
    return p


def add_number(doc: Document, text: str, *, bold_lead: str | None = None):
    return add_bullet(doc, text, bold_lead=bold_lead, style="List Number")


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int], *, header_fill: str = PALE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_text(cell, text, bold=True, color=NAVY, size=9.1)
    set_row_repeat(header)
    for row_values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for idx, text in enumerate(row_values):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_text(cell, text, bold=False, color=NAVY, size=9.0)
    set_table_geometry(table, widths)
    set_repeat_table_headers(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_three_cards(doc: Document, cards: Sequence[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    widths = [3120, 3120, 3120]
    set_table_geometry(table, widths)
    for idx, (title, body) in enumerate(cards):
        cell = table.cell(0, idx)
        set_cell_shading(cell, PALE_BLUE if idx != 1 else CALLOUT)
        set_cell_margins(cell, top=130, start=120, bottom=130, end=120)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "5", "color": LINE},
            bottom={"val": "single", "sz": "5", "color": LINE},
            start={"val": "single", "sz": "5", "color": LINE},
            end={"val": "single", "sz": "5", "color": LINE},
        )
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(title)
        set_run_font(r, size=10.2, color=BLUE_DARK, bold=True)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(body)
        set_run_font(r, size=9.2, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_source_table(doc: Document, sources: Sequence[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [2850, 6510]
    set_table_geometry(table, widths)
    for idx, text in enumerate(("Public material", "Link and use")):
        set_cell_shading(table.rows[0].cells[idx], PALE)
        set_cell_text(table.rows[0].cells[idx], text, bold=True, size=9.1)
    set_row_repeat(table.rows[0])
    for label, url, use in sources:
        row = table.add_row()
        set_row_cant_split(row)
        left, right = row.cells
        set_cell_width(left, widths[0])
        set_cell_width(right, widths[1])
        set_cell_margins(left)
        set_cell_margins(right)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_text(left, label, bold=True, size=9.0)
        right.text = ""
        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        display = url.removeprefix("https://")
        if len(display) > 52:
            display = "Open source ↗"
        else:
            display = display + " ↗"
        add_hyperlink(p, display, url)
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(use)
        set_run_font(r, size=8.8, color=SLATE)
    set_table_geometry(table, widths)
    set_repeat_table_headers(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def make_flow_diagram(path: Path) -> None:
    width, height = 1800, 390
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    font_regular = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 34)
    font_bold = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 37)
    font_small = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 27)

    nodes = [
        ("1", "Observation +\nlegal actions", "game offers choices"),
        ("2", "Policy selects\none offered ID", "reason + confidence"),
        ("3", "Validator +\nrunner", "canonical authority"),
        ("4", "Match runtime", "GameServer"),
        ("5", "Results + replay", "inspect decisions"),
    ]
    box_w, box_h = 292, 220
    left = 34
    gap = 72
    top = 70
    for idx, (number, title, subtitle) in enumerate(nodes):
        x0 = left + idx * (box_w + gap)
        x1 = x0 + box_w
        y0, y1 = top, top + box_h
        fill = PALE_BLUE if idx in (0, 1, 4) else CALLOUT
        draw.rounded_rectangle((x0, y0, x1, y1), radius=18, fill=f"#{fill}", outline=f"#{BLUE}", width=4)
        draw.ellipse((x0 + 18, y0 + 18, x0 + 66, y0 + 66), fill=f"#{BLUE}")
        number_bbox = draw.textbbox((0, 0), number, font=font_bold)
        nw = number_bbox[2] - number_bbox[0]
        nh = number_bbox[3] - number_bbox[1]
        draw.text((x0 + 42 - nw / 2, y0 + 42 - nh / 2 - 3), number, font=font_bold, fill=f"#{WHITE}")
        title_lines = title.split("\n")
        y = y0 + 82
        for line in title_lines:
            bbox = draw.textbbox((0, 0), line, font=font_bold)
            tw = bbox[2] - bbox[0]
            draw.text((x0 + (box_w - tw) / 2, y), line, font=font_bold, fill=f"#{NAVY}")
            y += 42
        bbox = draw.textbbox((0, 0), subtitle, font=font_small)
        tw = bbox[2] - bbox[0]
        draw.text((x0 + (box_w - tw) / 2, y1 - 42), subtitle, font=font_small, fill=f"#{SLATE}")
        if idx < len(nodes) - 1:
            arrow_y = y0 + box_h / 2
            start = x1 + 12
            end = x1 + gap - 12
            draw.line((start, arrow_y, end, arrow_y), fill=f"#{BLUE}", width=6)
            draw.polygon([(end, arrow_y), (end - 22, arrow_y - 14), (end - 22, arrow_y + 14)], fill=f"#{BLUE}")
    image.save(path, quality=95)


def add_picture_with_alt(doc: Document, path: Path, alt_text: str, width_inches: float = 6.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width_inches))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", "Proxy War decision flow")
    doc_pr.set("descr", alt_text)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def build() -> Path:
    WORK.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_flow_diagram(DIAGRAM)

    doc = Document()
    style_document(doc)
    section = doc.sections[0]
    configure_header_footer(section)
    set_first_page_header_footer(section)

    props = doc.core_properties
    props.title = "Let’s Play Proxy War — Guided Agent Lab Collaboration Brief"
    props.subject = "Existing Proxy War agent materials and a proposal for a guided notebook experience"
    props.author = "Proxy War"
    props.keywords = "Proxy War, Coworld, agent lab, notebook, replay, LegalAction"
    props.comments = "External-shareable collaboration brief assembled from public repository materials."

    # Cover
    p = doc.add_paragraph(style="PW Kicker")
    r = p.add_run("PROXY WAR  •  COLLABORATION BRIEF")
    set_run_font(r, size=9, color=BLUE, bold=True)
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_before = Pt(25)
    p.add_run("Let’s Play Proxy War")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("From a working starter policy to replay-driven improvement")
    set_paragraph_bottom_border(p, color=BLUE, size=10, space=10)

    add_callout(
        doc,
        "Recommendation",
        "Yes—the CrewRift learning format is the right shape for Proxy War. The hard infrastructure already exists. The opportunity is to connect it into one coherent, executable journey: understand a decision, change a policy, prove the choice is legal, run a match, inspect the replay, and improve one thing.",
        fill=PALE_BLUE,
    )

    add_three_cards(
        doc,
        [
            ("FOUNDATION", "Starter policies, a legal-action contract, packaging and upload tooling, the canonical match runtime, replays, and decision evidence."),
            ("MISSING LAYER", "A guided experience that makes the whole path visible, runnable, and understandable to someone new to Proxy War."),
            ("THE ASK", "Adapt the CrewRift play-and-analyze loop as a learning layer over the existing system—not as a new SDK or game implementation."),
        ],
    )

    metadata = doc.add_table(rows=3, cols=2)
    set_table_geometry(metadata, [2050, 7310])
    metadata_rows = [
        ("PURPOSE", "Align on the desired notebook experience and the materials already available."),
        ("AUDIENCE", "A technical learning-experience collaborator familiar with notebooks, containers, and agents; no prior Proxy War context assumed."),
        ("STATUS", "External-shareable working brief • Stable protocol facts separated from live service details."),
    ]
    for i, (label, value) in enumerate(metadata_rows):
        left, right = metadata.rows[i].cells
        set_cell_shading(left, PALE)
        set_cell_shading(right, WHITE)
        set_cell_text(left, label, bold=True, color=BLUE_DARK, size=8.5)
        set_cell_text(right, value, size=9.1)
    set_table_geometry(metadata, [2050, 7310])

    add_page_break(doc)

    # Page 2
    add_label(doc, "context and recommendation", proposed=True)
    doc.add_heading("1. The opportunity", level=1)
    add_body(
        doc,
        "The useful part of CrewRift’s “Let’s Play” notebook is its learning loop: contract → working policy → offline examples → packaged agent → real game → replay → one controlled improvement. Proxy War already has each technical piece. What it lacks is one polished path across them.",
        keep=True,
    )
    add_body(
        doc,
        "The notebook should expose and orchestrate the system we have. It should not generate a parallel client, rewrite the protocol in Python, or introduce a second validator or runner.",
        keep=True,
    )
    add_three_cards(
        doc,
        [
            ("KEEP", "Progressive explanation, editable examples, frozen scenarios, a clean run, replay analysis, and a clear next experiment."),
            ("ADAPT", "Replace CrewRift-specific actions and fixtures with Proxy War observations, offered legal actions, replays, and telemetry."),
            ("AVOID", "Hard-coded league IDs, unpinned environments, notebook-only protocol code, hidden fallbacks, or publishing from “Run All.”"),
        ],
    )

    add_label(doc, "product and contract")
    doc.add_heading("2. Proxy War in 60 seconds", level=1)
    add_body(
        doc,
        "Proxy War is an autonomous strategy game in which software agents claim territory, build an economy, form alliances, attack rivals, and play complete matches without human input. An agent—called a policy—runs in its own container. At each decision, it receives the current observation and a list of actions that are legal now.",
        keep=True,
    )
    add_picture_with_alt(
        doc,
        DIAGRAM,
        "Five-step flow: the game sends an observation and legal actions; the policy selects one offered action ID; the canonical validator and runner check and execute it; the GameServer runs the match; results and a replay are produced for inspection.",
    )
    add_callout(
        doc,
        "The one rule that matters",
        "A policy selects one exact LegalAction.id from the list it was offered. It never constructs or sends a raw game-engine intent. An invented or stale ID is rejected by the existing validator.",
        fill=PALE_BLUE,
    )

    add_page_break(doc)

    # Page 3
    add_label(doc, "public repository inventory")
    doc.add_heading("3. What already exists", level=1)
    add_body(doc, "The material is substantial. The missing piece is a guided experience across it—not another protocol or starter implementation.", keep=True)
    add_table(
        doc,
        ("Asset", "What it provides", "Notebook role"),
        [
            ("Starter README + onboarding", "Public landing page and first-policy walkthrough.", "Orient the learner; link out for setup details."),
            ("llm-player.mjs", "Editable LLM policy with background planning and guarded action selection.", "Expose the three safe customization points."),
            ("starter-player.mjs", "Compact non-LLM baseline.", "Supply a deterministic control policy."),
            ("launch.sh", "Preflight, container build, browser sign-in, upload, and policy-version ID lookup.", "Keep upload as an explicit optional terminal step."),
            ("Player protocol", "Minimal decision request/response contract.", "Anchor the contract tour and fixtures; use the starter for full transport behavior."),
            ("Adapter certification + episode tooling", "Canonical local integration and replay-verification paths.", "Invoke behind the notebook; do not reimplement."),
            ("Replays + decision evidence", "Rendered match evidence plus policy-authored reason summaries, confidence, and degradation signals.", "Power the analyze-and-improve chapter from a reviewed bundle or authorized view."),
        ],
        (2050, 3860, 3450),
    )
    add_callout(
        doc,
        "Two layers, one architecture",
        "The lightweight external starter is the learner-facing policy. The Coworld adapter is the engine/integration layer. The notebook may connect them, but it must keep their responsibilities distinct.",
    )

    add_label(doc, "starter behavior")
    doc.add_heading("4. How the supplied LLM starter works", level=1)
    add_body(
        doc,
        "The starter separates slow strategic reasoning from time-sensitive move selection. A model writes a compact standing plan in the background; the decision handler immediately maps the latest good plan to one currently offered legal action.",
        keep=True,
    )
    add_table(
        doc,
        ("Builder edits", "Purpose", "Boundary"),
        [
            ("STRATEGY", "Plain-language standing orders: how the nation should play.", "Changes goals, not the transport contract."),
            ("buildState", "Selects the game facts shown to the planner.", "Uses the observation; does not bypass it."),
            ("choose", "Maps the standing plan to one offered action.", "Must return an item from legalActions."),
        ],
        (2100, 3860, 3400),
    )
    add_body(
        doc,
        "The starter also keeps anti-loop memory, sanitizes untrusted game text, continues from the last good plan when planning fails, and emits fallbackUsed and llmPlannerDegraded so a strategic loss can be distinguished from an unhealthy planner.",
        keep=True,
    )

    # Page 4
    add_label(doc, "current workflow")
    doc.add_heading("5. The builder journey today", level=1)
    add_number(doc, "Check the environment. Run bash launch.sh --doctor for a non-mutating preflight.", bold_lead="Check the environment.")
    add_number(doc, "Choose a starting point. Use the LLM policy or the compact rule-based baseline.", bold_lead="Choose a starting point.")
    add_number(doc, "Customize the policy. Edit STRATEGY, buildState, and choose while preserving the transport and legal-action contract.", bold_lead="Customize the policy.")
    add_number(doc, "Validate through canonical tooling. Exercise fixtures and the existing adapter’s certification or verified episode path.", bold_lead="Validate through canonical tooling.")
    add_number(doc, "Build and upload deliberately. The launcher can package and upload a policy after browser authentication.", bold_lead="Build and upload deliberately.")
    add_number(doc, "Seat or submit the policy. A separate invited-match or current league step is required.", bold_lead="Seat or submit the policy.")
    add_number(doc, "Inspect and improve. Use the result, rendered replay, decision reasons, confidence, and degradation telemetry.", bold_lead="Inspect and improve.")

    add_callout(
        doc,
        "Important correction to the current onboarding language",
        "bash launch.sh my-agent builds and uploads a policy version. It does not itself run a match. That version must then be seated in an invited episode or submitted to a live league.",
        fill="FFF7E8",
        title_color="A85D00",
    )

    doc.add_heading("The current starter commands", level=2)
    add_code_block(
        doc,
        "git clone https://github.com/0xNad/ProxyWar.git\n"
        "cd ProxyWar/coworld-adapter/tester-starter-llm\n\n"
        "# Non-mutating environment check\n"
        "bash launch.sh --doctor\n\n"
        "# Explicit authenticated build + upload\n"
        "bash launch.sh my-agent",
    )
    add_body(
        doc,
        "Hosted authentication, model access, league identifiers, and current CLI compatibility are live-service facts. Recheck them at the time of use; do not freeze them into the durable notebook.",
        after=0,
        keep=True,
    )

    add_page_break(doc)

    # Page 5
    add_label(doc, "desired learning experience", proposed=True)
    doc.add_heading("6. Desired notebook experience", level=1)
    add_body(
        doc,
        "Present this as one journey with two linked chapters. They may live in one notebook or two, but the learner should experience a continuous loop from a working decision to a measured improvement.",
        keep=True,
    )
    modules = doc.add_table(rows=1, cols=2)
    set_table_geometry(modules, [4680, 4680])
    left, right = modules.rows[0].cells
    for cell, fill in ((left, PALE_BLUE), (right, CALLOUT)):
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=150, start=120, bottom=150, end=120)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "5", "color": LINE},
            bottom={"val": "single", "sz": "5", "color": LINE},
            start={"val": "single", "sz": "5", "color": LINE},
            end={"val": "single", "sz": "5", "color": LINE},
        )
    left.text = ""
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("CHAPTER A")
    set_run_font(r, size=8, color=BLUE, bold=True)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Let’s Play Proxy War")
    set_run_font(r, size=13, color=BLUE_DARK, bold=True)
    for text in (
        "Preview a rendered replay before setup.",
        "Inspect one frozen decision request.",
        "Run the rule baseline on that fixture.",
        "Edit strategy, state selection, or choice logic.",
        "Prove the selected ID was offered.",
        "Invoke a canonical episode and verify its replay.",
        "Open the evidence produced by the match.",
    ):
        p = left.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.27)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        set_run_font(r, size=9.4)

    right.text = ""
    p = right.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("CHAPTER B")
    set_run_font(r, size=8, color=BLUE, bold=True)
    p = right.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Let’s Analyze & Improve")
    set_run_font(r, size=13, color=BLUE_DARK, bold=True)
    for text in (
        "Read the result and territory trajectory.",
        "Summarize action mix and decision reasons.",
        "Separate strategy failure from fallback or degradation.",
        "Form one explicit hypothesis.",
        "Change one behavior—not five at once.",
        "Rerun the same controlled scenario.",
        "Compare baseline and challenger evidence.",
    ):
        p = right.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.27)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        set_run_font(r, size=9.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    add_callout(
        doc,
        "Safe “Run All” behavior",
        "The core notebook path must be deterministic and non-publishing. Upload, league submission, browser sign-in, and any other hosted mutation belong in a clearly separated optional final section that requires explicit user action.",
        fill=PALE_BLUE,
    )

    add_label(doc, "architecture constraints", proposed=True)
    doc.add_heading("7. Guardrails", level=1)
    add_table(
        doc,
        ("Do", "Do not"),
        [
            ("Invoke the existing Node policy and canonical adapter paths.", "Rewrite the protocol or policy runtime as notebook-only Python."),
            ("Choose exactly one supplied LegalAction.id.", "Construct raw OpenFront intents or bypass validation."),
            ("Use frozen, sanitized fixtures for offline teaching.", "Depend on a changing hosted league for the core lesson."),
            ("Expose fallback and planner-degradation telemetry.", "Hide an unhealthy strategic brain behind legal executor output."),
            ("Fail loudly with a useful next step.", "Save secrets, private paths, or raw internal logs in outputs."),
        ],
        (4680, 4680),
    )

    add_page_break(doc)

    # Page 6
    add_label(doc, "scope and quality bar", proposed=True)
    doc.add_heading("8. Deliverables", level=1)
    add_bullet(doc, "A guided “Let’s Play Proxy War” chapter and a linked “Let’s Analyze & Improve” chapter; one notebook or two is acceptable.")
    add_bullet(doc, "A small frozen and sanitized evidence bundle: representative decision requests, one result, one verified replay, and decision metadata.")
    add_bullet(doc, "A minimal helper layer that calls canonical policy and adapter code instead of duplicating it.")
    add_bullet(doc, "A concise README and a pinned or recorded tested environment.")
    add_bullet(doc, "Executed HTML or PDF exports with meaningful outputs, plus an automated clean-run check.")
    add_bullet(doc, "Existing licenses and attribution retained for every reused component.")

    doc.add_heading("9. Definition of done", level=1)
    acceptance = [
        "A fresh macOS, Linux, or WSL user can run the core path top to bottom without internal-repository access or hosted credentials.",
        "Every saved executable cell has an execution count and meaningful output; a clean-room or CI run proves it.",
        "The selected action is always one of the supplied legal IDs, and an invalid-ID fixture visibly fails.",
        "A full episode uses the canonical adapter, validator, and runner; the resulting replay verifies.",
        "Analysis shows the outcome, action mix, decision reasons, and fallback/planner-degradation signals.",
        "The learner completes one reproducible baseline-versus-change comparison.",
        "“Run All” cannot upload, submit, publish, or trigger browser authentication.",
        "No secrets, private paths, internal project-state files, or unreviewed raw logs are included.",
        "Errors fail loudly and point to the next concrete action.",
    ]
    for item in acceptance:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("☐  ")
        set_run_font(r, size=10.2, color=BLUE, bold=True)
        r = p.add_run(item)
        set_run_font(r, size=10.1)

    add_callout(
        doc,
        "Success in one sentence",
        "A successful first version takes someone with no Proxy War context from a clean checkout to a validated agent decision and a replay-based improvement experiment, with every step executed, inspectable, and reproducible.",
        fill=PALE_BLUE,
    )

    add_page_break(doc)

    # Page 7
    add_label(doc, "external handoff", proposed=True)
    doc.add_heading("10. Public-safe handoff", level=1)
    add_body(
        doc,
        "Share these curated public links. Do not send the whole development workspace: it adds internal operating material without helping the collaborator understand the agent path.",
        keep=True,
    )
    add_source_table(
        doc,
        [
            ("Current Proxy War source", "https://github.com/0xNad/ProxyWar", "Moving public landing page for the current codebase."),
            ("Reviewed starter snapshot", "https://github.com/0xNad/ProxyWar/tree/333bdc39a3e6d2653bef208f46b100b4b7d18e5c/coworld-adapter/tester-starter-llm", "Immutable technical source used for this brief. Treat hosted-runtime language inside it as dated and recheck it."),
            ("Reviewed protocol snapshot", "https://github.com/0xNad/ProxyWar/blob/333bdc39a3e6d2653bef208f46b100b4b7d18e5c/coworld-adapter/docs/player-protocol.md", "Immutable minimal decision contract; the starter shows the wider transport lifecycle."),
            ("Coworld source", "https://github.com/Metta-AI/coworld", "The upstream packaging and league tooling."),
            ("Public league view", "https://beta.proxywar.xyz/league", "Live demonstration; treat current contents as changeable, not source truth."),
            ("Softmax Observatory", "https://softmax.com/observatory", "Hosted policy and match evidence; Softmax sign-in is currently required."),
        ],
    )

    doc.add_heading("Materials still to provide for the prototype", level=2)
    add_bullet(doc, "One scrubbed decision_request fixture that exercises a meaningful choice.")
    add_bullet(doc, "One reviewed replay/results bundle with decision metadata and no private logs.")
    add_bullet(doc, "One stable replay screenshot or public-safe preview asset.")
    add_bullet(doc, "The exact environment versions proven by the clean-run check.")

    add_callout(
        doc,
        "Source synchronization note",
        "At the time of this audit, the separate proxywar-coworld-starter repository trails the newer starter inside the main ProxyWar repository. Use the reviewed technical snapshot above until the standalone repository is synchronized, and recheck any hosted-service claims inside it.",
        fill="FFF7E8",
        title_color="A85D00",
    )

    doc.add_heading("Keep out of the handoff", level=2)
    add_body(
        doc,
        "Internal project-state, operating-layer files, competitive strategy documents, raw model prompts or logs, temporary certification artifacts, credentials, authentication material, and operator runbooks.",
        keep=True,
    )

    add_label(doc, "first milestone", proposed=True)
    doc.add_heading("11. Recommended prototype", level=1)
    add_callout(
        doc,
        "Prove the learning loop before the full hosted path",
        "Frozen request → starter decision → canonical legal-ID validation → supplied replay/results bundle → one diagnostic chart. Once that is clean, add a full local episode and keep hosted upload/submission as an explicit optional step.",
        fill=PALE_BLUE,
    )
    add_body(
        doc,
        "This is deliberately smaller than a complete end-to-end integration. It proves the collaborator can teach the right architecture and create a useful analysis loop before either side spends time on live-service edge cases.",
        after=6,
        keep=True,
    )

    doc.add_heading("Working split for the first prototype", level=2)
    add_table(
        doc,
        ("Proxy War provides", "Notebook collaborator provides"),
        [
            ("Reviewed source and protocol snapshots.", "An executed, progressive notebook structure."),
            ("A scrubbed decision-request fixture.", "A minimal helper layer and clear visual outputs."),
            ("A reviewed replay/results evidence bundle.", "One baseline-versus-change analysis experiment."),
            ("Canonical validation hooks and acceptance review.", "Recorded versions, clean-run proof, and an export."),
        ],
        (4680, 4680),
    )
    add_callout(
        doc,
        "Decision requested",
        "Confirm whether this play → analyze → improve shape matches the collaboration you had in mind. If it does, the next step is to exchange the scrubbed fixture and evidence bundle and build the deliberately small prototype above.",
        fill=PALE_BLUE,
    )

    # Force update fields on open.
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(DRAFT)
    return DRAFT


if __name__ == "__main__":
    print(build())
