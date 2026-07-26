from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from build_lets_play_proxywar_brief import (
    BLUE,
    BLUE_DARK,
    CALLOUT,
    GREEN,
    LINE,
    MUTED,
    NAVY,
    PALE,
    PALE_BLUE,
    SLATE,
    WHITE,
    add_body,
    add_bullet,
    add_callout,
    add_label,
    add_page_break,
    add_page_number,
    add_picture_with_alt,
    add_source_table,
    add_table,
    add_three_cards,
    make_flow_diagram,
    rgb,
    set_first_page_header_footer,
    set_paragraph_bottom_border,
    set_run_font,
    style_document,
)


ROOT = Path("/Users/claude/Documents/proxywar_main")
WORK = ROOT / "tmp" / "proxywar_notebook_context"
PACKAGE = ROOT / "outputs" / "share" / "ProxyWar_Notebook_Context_Packet"
DRAFT = WORK / "ProxyWar_Notebook_Context_draft.docx"
FINAL = PACKAGE / "ProxyWar_Notebook_Context.docx"
DIAGRAM = WORK / "proxywar_decision_flow.png"


def configure_context_header_footer(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "PROXY WAR  /  NOTEBOOK CONTEXT"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.runs[0], size=8.5, color=MUTED, bold=True)
    set_paragraph_bottom_border(paragraph, color=LINE)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    add_page_number(paragraph)

    section.different_first_page_header_footer = True
    first_header = section.first_page_header
    paragraph = first_header.paragraphs[0]
    paragraph.text = "PROXY WAR"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.runs[0], size=8.5, color=BLUE, bold=True)
    set_paragraph_bottom_border(paragraph, color=BLUE, size=8)

    first_footer = section.first_page_footer
    paragraph = first_footer.paragraphs[0]
    paragraph.text = "NOTEBOOK CONTEXT  •  JULY 2026"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    set_run_font(paragraph.runs[0], size=8.2, color=MUTED, bold=True)


def build() -> Path:
    WORK.mkdir(parents=True, exist_ok=True)
    PACKAGE.mkdir(parents=True, exist_ok=True)
    make_flow_diagram(DIAGRAM)

    doc = Document()
    style_document(doc)
    section = doc.sections[0]
    configure_context_header_footer(section)

    properties = doc.core_properties
    properties.title = "Proxy War Notebook Context"
    properties.subject = "A concise, exploratory overview of existing Proxy War agent materials"
    properties.author = "Proxy War"
    properties.keywords = "Proxy War, Coworld, agent notebook, replay, LegalAction"
    properties.comments = "External-shareable context note assembled from public repository materials."

    # Page 1: neutral opening.
    paragraph = doc.add_paragraph(style="PW Kicker")
    run = paragraph.add_run("PROXY WAR  •  NOTEBOOK CONTEXT")
    set_run_font(run, size=9, color=BLUE, bold=True)

    paragraph = doc.add_paragraph(style="Title")
    paragraph.paragraph_format.space_before = Pt(25)
    paragraph.add_run("Could “Let’s Play” fit Proxy War?")

    paragraph = doc.add_paragraph(style="Subtitle")
    paragraph.add_run("A concise view of the existing agent path and why the format feels relevant")
    set_paragraph_bottom_border(paragraph, color=BLUE, size=10, space=10)

    add_callout(
        doc,
        "Short answer",
        "Yes—the guided format looks relevant to Proxy War. The material in this packet is shared only as context for that reaction. It is not a specification, assignment, or assumption that you want to build it.",
        fill=PALE_BLUE,
    )

    add_three_cards(
        doc,
        [
            (
                "WHAT EXISTS",
                "A Coworld starter policy, legal-action and replay contracts, canonical validation, and match evidence.",
            ),
            (
                "WHY IT MAPS",
                "The CrewRift flow—understand, edit, test, run, replay, improve—matches the natural Proxy War agent-learning loop.",
            ),
            (
                "WHAT IS OPEN",
                "Whether this is interesting, what form it might take, and whether any follow-up is useful are deliberately left open.",
            ),
        ],
    )

    add_body(
        doc,
        "This packet contains a short overview plus a few deliberately selected public examples. Operational commands, live league identifiers, internal project notes, and time-sensitive hosted claims are intentionally excluded.",
        keep=True,
    )

    add_page_break(doc)

    # Page 2: stable product and architecture context.
    add_label(doc, "stable product context")
    doc.add_heading("1. Proxy War in 60 seconds", level=1)
    add_body(
        doc,
        "Proxy War is an autonomous strategy game in which software agents claim territory, build an economy, form alliances, attack rivals, and play complete matches without human input. At each decision, the game gives a policy the current observation and a list of actions that are legal at that moment.",
        keep=True,
    )
    add_picture_with_alt(
        doc,
        DIAGRAM,
        "Five-step flow: the game sends an observation and legal actions; the policy selects one offered action ID; the canonical validator and runner check and execute it; the GameServer runs the match; results and a replay are produced for inspection.",
    )
    add_callout(
        doc,
        "The durable contract",
        "A policy returns one exact LegalAction.id from the list it was offered. It does not construct a raw game-engine intent. The existing validator, runner, and GameServer remain authoritative.",
        fill=PALE_BLUE,
    )

    doc.add_heading("2. What already exists", level=1)
    add_table(
        doc,
        ("Material", "What it contributes", "Status in this packet"),
        [
            (
                "Coworld starter policy",
                "A small rule-based agent with one clearly marked strategy function.",
                "Included as the simplest working reference.",
            ),
            (
                "Player protocol",
                "The minimal decision_request → decision_response contract.",
                "Included as a reviewed reference copy.",
            ),
            (
                "Global and replay protocol",
                "The status, spectator, result, and saved-replay surfaces.",
                "Included without time-sensitive league details.",
            ),
            (
                "Canonical match path",
                "Validation, episode execution, results, and replay artifacts.",
                "Explained here; live hosted details intentionally omitted.",
            ),
        ],
        (2300, 3830, 3230),
    )

    add_callout(
        doc,
        "One architecture, not a notebook-only SDK",
        "Any notebook-shaped experience would be most useful if it exposed the existing policy and match path. It should not introduce a second protocol, validator, runner, or raw-intent route.",
        fill=CALLOUT,
    )

    add_page_break(doc)

    # Page 3: exploratory mapping, contents, and safe links.
    add_label(doc, "exploratory mapping", proposed=True)
    doc.add_heading("3. How the format might translate", level=1)
    add_body(
        doc,
        "The strongest overlap is the learning sequence, not any specific CrewRift implementation detail. A Proxy War version could, in principle:",
        keep=True,
    )
    add_bullet(doc, "Show a replay or match outcome first, so the learner understands the destination.")
    add_bullet(doc, "Open one frozen decision request and explain the observation plus offered legal actions.")
    add_bullet(doc, "Run a working baseline policy, then change one small piece of strategy or selection logic.")
    add_bullet(doc, "Verify that the policy selected one offered LegalAction.id through the canonical path.")
    add_bullet(doc, "Inspect replay/result evidence and compare one controlled change with the baseline.")

    add_callout(
        doc,
        "No implied ask",
        "This is simply the part of the CrewRift notebook that appears transferable. It is not a proposed scope, commitment, or request for you to take ownership of it.",
        fill=PALE_BLUE,
    )

    doc.add_heading("4. What is included", level=1)
    add_table(
        doc,
        ("File", "Purpose"),
        [
            ("README_FIRST.md", "Packet orientation and current-use caveats."),
            ("SEND_WITH_THIS.txt", "A short, neutral message to accompany the packet."),
            ("ProxyWar_Notebook_Context.docx", "This concise overview."),
            ("protocol/", "The player and global/replay contracts."),
            ("reference-agent/", "A small non-LLM Coworld policy baseline."),
            ("licenses/", "Applicable license copies for the selected source material."),
        ],
        (3220, 6140),
    )

    doc.add_heading("5. Public source links", level=1)
    add_source_table(
        doc,
        [
            (
                "Proxy War source",
                "https://github.com/0xNad/ProxyWar",
                "Platform and protocol source of truth.",
            ),
            (
                "Proxy War Coworld adapter",
                "https://github.com/0xNad/ProxyWar/tree/main/coworld-adapter",
                "Active Coworld integration and the source of the selected examples.",
            ),
            (
                "Coworld source",
                "https://github.com/Metta-AI/coworld",
                "Upstream packaging and league tooling.",
            ),
        ],
    )
    # The shared source-table helper leaves an empty trailing paragraph. On this
    # compact final page, LibreOffice places that paragraph on a blank page.
    trailing = doc.paragraphs[-1]
    if not trailing.text.strip():
        trailing._element.getparent().remove(trailing._element)

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(DRAFT)
    return DRAFT


if __name__ == "__main__":
    print(build())
